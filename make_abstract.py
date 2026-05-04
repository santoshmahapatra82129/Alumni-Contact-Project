from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins similar to reference
for section in doc.sections:
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ABSTRACT")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

doc.add_paragraph()

def add_para(text, italic=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic
    return p

add_para(
    "The Alumni Contact project is a web-based networking and career-guidance "
    "platform designed to bridge the communication gap between current students "
    "and graduated alumni of an engineering institution. The platform centralises "
    "company-wise placement experiences, interview insights, and learning "
    "resources, enabling students to prepare effectively for campus recruitment "
    "while empowering alumni to share their professional journeys with juniors."
)
add_para(
    "The system is developed using Node.js and Express.js on the server side, "
    "with EJS templating for dynamic views and MongoDB (via Mongoose) as the "
    "primary data store. Authentication is handled through Passport.js with "
    "Passport-Local-Mongoose, supporting role-based registration and login for "
    "two distinct user types — students and alumni — each redirected to a "
    "tailored profile dashboard upon successful sign-in. Session management is "
    "implemented through Express-Session with secure HTTP-only cookies."
)
add_para(
    "Core modules of the platform include a Companies directory where students "
    "can browse recruiting organisations and read approved interview experiences "
    "contributed by seniors, an Experience submission portal capturing rounds, "
    "questions, tips and difficulty levels, a Mock Test engine that fetches "
    "questions from the database and auto-evaluates the score, a Resources hub "
    "covering aptitude, DSA, mock interviews and downloadable LaTeX resume "
    "templates, and role-specific profile pages for students (CGPA, branch, "
    "batch) and alumni (company, designation, LinkedIn)."
)
add_para(
    "The project demonstrates the practical application of full-stack web "
    "development concepts including RESTful routing, MVC architecture, "
    "schema-based data modelling, middleware-driven authentication, automated "
    "database seeding, and graceful fallback to an in-memory MongoDB instance "
    "for portable deployment."
)
add_para(
    "To ensure a secure and trustworthy academic environment, the platform "
    "enforces role-based access control, protects sensitive routes through "
    "authentication middleware, and validates user identity against the "
    "registered role at every login attempt. The final outcome is a scalable, "
    "extensible alumni-engagement framework that fosters mentorship, knowledge "
    "transfer and placement readiness within the institution.",
    italic=True
)

out = r"D:\MAjor\Alumini-Contact-Minor-Project-main (1)\Alumini-Contact-Minor-Project-main\Abstract.docx"
doc.save(out)
print(out)
